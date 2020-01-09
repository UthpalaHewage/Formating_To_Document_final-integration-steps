from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document()


def write_to_doc_method(content_dict):
    for key in content_dict:
        if str(key) == "Main Topic":
            document.add_heading(str(content_dict[key][0].content).upper(), 0)
        else:
            previous_title = ""
            for value in content_dict[key]:

                new_title = value.title

                if new_title != previous_title and value.title_availability:
                    document.add_heading(value.title, level=1)
                    previous_title = new_title

                visual_content = document.add_paragraph()
                img_content = document.add_paragraph()
                run_image_content = img_content.add_run()

                if value.content_availability:
                    visual_content.add_run().add_break()
                    run_visual_content = visual_content.add_run(value.content)
                    run_visual_content.bold = True
                    run_visual_content.italic = True
                    if value.figure:
                        run_image_content.add_picture('figures/' + str(value.frame_position) + ".jpg", width=Inches(5),
                                                      height=Inches(3))
                        img_content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    if value.figure:
                        run_image_content.add_picture('figures/' + str(value.frame_position) + ".jpg", width=Inches(5),
                                                      height=Inches(3))
                        img_content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                visual_content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                transcript_content = document.add_paragraph()
                transcript_content.add_run("Hansi's transcript content")
                transcript_content.add_run("\n")

    document.save('demo.docx')
